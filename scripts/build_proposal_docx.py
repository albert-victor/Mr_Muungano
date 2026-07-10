"""Generate MuunganoAi proposal and form answers as .docx for submission."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT_PROPOSAL = ROOT / "PROPOSAL_MuunganoAi.docx"
OUT_SUBMISSION = ROOT / "SUBMISSION_MuunganoAi.docx"

FORM_ANSWERS: list[tuple[str, str]] = [
    (
        "Jina la Mradi",
        "MuunganoAi — Msaidizi wa AI wa Elimu ya Muungano wa Tanzania",
    ),
    (
        "Muhtasari mfupi",
        "MuunganoAi ni msaidizi wa mazungumzo unaosaidia Watanzania — hasa vijana — "
        "kuelewa Muungano wa Tanganyika na Zanzibar. Majibu yanatokana na nyaraka rasmi "
        "(Katiba, Articles of Union, machapisho ya serikali), si kutoka kwa mawazo ya AI "
        "peke yake. Mtumiaji anaweza kuuliza kwa maandishi au sauti, na kupokea jibu kwa "
        "njia ile ile anayopendelea.",
    ),
    (
        "Tatizo",
        "Wengi hawajui kwa undani Muungano ulivyoanzishwa, kwa nini Tanzania ina muundo wa "
        "serikali mbili, na “Mambo ya Muungano” ni nini. Nyaraka rasmi ni ndefu na ngumu; "
        "mtandaoni kuna mchanganyiko wa ukweli na maoni. Hii inapunguza uelewa wa uraia, "
        "hasa kwa kizazi kinachokua na kidijitali.",
    ),
    (
        "Suluhisho",
        "Tumejenga mfumo wa RAG (Retrieval-Augmented Generation) unaofanya hivi:\n"
        "1. Mtumiaji anauliza swali — kwa kuandika au kuzungumza\n"
        "2. Mfumo unatafuta taarifa sahihi kutoka maktaba ya nyaraka zilizothibitishwa\n"
        "3. AI inaandika jibu kulingana na vyanzo vilivyopatikana tu\n"
        "4. Jibu linapatikana kwa maandishi; mtumiaji anaweza pia kulisikiliza kwa sauti\n\n"
        "Mfumo unasaidia Kiswahili na Kiingereza, na mitindo mbalimbali ya majibu (vijana, "
        "mwanafunzi, rasmi). Vyanzo visivyothibitishwa havijumuishwi — tunatumia allowlist "
        "ya tovuti na nyaraka rasmi.",
    ),
    (
        "Je, AI inatumiaje?",
        "Semantic search — kutafuta vipande vinavyohusiana na swali (embeddings za lugha mbili)\n"
        "Large Language Model — kuandika majibu ya mazungumzo kutoka muktadha uliopatikana\n"
        "Speech-to-Text / Text-to-Speech — kuruhusu mazungumzo ya kawaida (andika au zungumza, "
        "soma au sikiliza)\n"
        "Prioritization — Katiba na Articles of Union zina uzito mkubwa kuliko vyanzo vingine",
    ),
    (
        "Lengo / Watumiaji",
        "Vijana, wanafunzi, walimu, na raia wanaotaka kuelewa historia, muundo, na sheria za "
        "Muungano — bila kusoma mamia ya kurasa za PDF.",
    ),
    (
        "Athari kwa Maendeleo",
        "MuunganoAi inachangia elimu ya uraia na uelewa wa taasisi za kitaifa — msingi wa "
        "ushiriki wenye maarifa katika jamii ya kidijitali. Inaendana na mada “AI kwa Maendeleo” "
        "kwa kutoa elimu sahihi, inayopatikana, na yenye msingi wa vyanzo rasmi.",
    ),
    (
        "Ubunifu",
        "Tofauti na chatbot za jumla, MuunganoAi ni msaidizi maalum wa kitaifa unaojikita kwenye "
        "Muungano wa Tanzania. Inaunganisha RAG, vyanzo vilivyochaguliwa kwa makini, lugha mbili, "
        "na uzoefu wa mazungumzo (maandishi na sauti) — ili elimu iwe rahisi kufikia na kuaminika.",
    ),
    (
        "Tech Stack",
        "Python · FastAPI · ChromaDB · sentence-transformers · OpenRouter (Gemini) · "
        "Frontend (HTML/CSS/JS) · Web Speech APIs kwa sauti",
    ),
    (
        "Hali ya Mradi",
        "Mradi umekamilika takriban 90%. Mfumo wa msingi unafanya kazi: ingest ya nyaraka, "
        "utafutaji wa semantic, chat na streaming, lugha mbili, na uwezo wa sauti (ingizo na "
        "majibu). Maktaba ya maarifa ina mamia ya vipande vya nyaraka rasmi. Kazi zilizobaki ni "
        "uboreshaji mdogo wa UI na upanuzi wa vyanzo.\n\n"
        "Demo: python run.py → http://127.0.0.1:8000\n"
        "[Jaza: GitHub / demo URL]",
    ),
    (
        "Kwa nini mradi huu?",
        "Kwa sababu unashughulikia tatizo halisi la elimu ya uraia — si teknolojia kwa ajili ya "
        "teknolojia. MuunganoAi inaonyesha jinsi AI inavyoweza kutumika kwa maendeleo: kutoa "
        "ukweli kutoka vyanzo vinavyothibitishwa, kwa lugha rahisi, kwa watu wengi zaidi.",
    ),
    (
        "Mfumo wako utawafikiaje vijana?",
        "Vijana hawataki PDF ndefu — wanataka majibu ya haraka, kwa lugha wanayoielewa, kwenye "
        "simu zao. MuunganoAi imejengwa kwa ajili hiyo.\n\n"
        "Tutawafikia kwanza mtandaoni: app inayofunguka moja kwa moja kwenye simu au kompyuta, "
        "bila kusakinisha chochote. Anaweza kuuliza “Muungano ulianzishwa lini?” au “Kwa nini "
        "serikali mbili?” na kupata jibu la dakika moja — kwa Kiswahili au Kiingereza, kwa mtindo "
        "wa vijana au mwanafunzi.\n\n"
        "Pili, mitandao ya kijamii — Instagram, TikTok, WhatsApp. Tutashiriki maswali ya kawaida "
        "na majibu mafupi yanayovuta, halafu kuwaelekeza kwenye app kwa kujifunza zaidi.\n\n"
        "Tatu, shuleni na vyuo — MuunganoAi ni msaidizi wa walimu wa Historia na Uraia, si mbadala "
        "wa mtaala. Mwanafunzi anaweza kuuliza maswali baada ya somo au kujiandaa kwa mitihani.\n\n"
        "Na kwa kuwa tuko kwenye mazingira ya TEHAMA kama Girls in ICT, tutaonyesha kwa vitendo "
        "kwamba AI inaweza kuwa zana ya elimu ya uraia, si burudani tu.\n\n"
        "Kwa ufupi: tunakwenda pale vijana walipo — simu, mitandao, darasa — na tunazungumza lugha yao.",
    ),
    (
        "Mfumo huu unatumia teknolojia gani?",
        "MuunganoAi inatumia RAG (Retrieval-Augmented Generation) — teknolojia inayohakikisha AI "
        "haijibu kutoka kwa mawazo yake peke yake, bali kutoka nyaraka rasmi zilizohifadhiwa kwenye "
        "maktaba ya kidijitali.\n\n"
        "Backend: Python, FastAPI\n"
        "Hifadhi ya maarifa: ChromaDB (vector database)\n"
        "Utafutaji: sentence-transformers (paraphrase-multilingual-MiniLM-L12-v2) — inaelewa "
        "Kiswahili na Kiingereza\n"
        "AI ya mazungumzo: Gemini 2.5 Flash kupitia OpenRouter\n"
        "Frontend: HTML, CSS, JavaScript — inafunguka kwenye simu na kompyuta bila kusakinisha\n"
        "Sauti: Web Speech API (STT/TTS) — sehemu ya kawaida ya mazungumzo, si lengo kuu\n\n"
        "Mfumo pia una crawler wa tovuti rasmi (allowlist) na ingest ya PDF, ili maktaba ya "
        "maarifa ibaki sahihi na ya kisasa.",
    ),
    (
        "Mfumo wako utasaidiaje kuongeza uelewa wa Muungano?",
        "Badala ya kusoma mamia ya kurasa za Katiba, vijana na raia watauliza swali moja — mfano: "
        "“Muungano ulianzishwa lini?”, “Mambo ya Muungano ni nini?”, “Kwa nini kuna serikali mbili?” "
        "— na kupata jibu linaloeleweka, kwa lugha wanayopendelea.\n\n"
        "MuunganoAi inabadilisha nyaraka ngumu kuwa mazungumzo ya kawaida, yenye msingi wa vyanzo "
        "vilivyothibitishwa. Mtumiaji anachagua mtindo (vijana, mwanafunzi, rasmi) na lugha (SW/EN), "
        "hivyo elimu inafika mtu alivyo.\n\n"
        "Kwa kuwa majibu yanatokana na Katiba, Articles of Union, na machapisho ya serikali, "
        "tunapunguza uvumi na mchanganyiko wa taarifa mtandaoni. Hii inajenga uraia wenye maarifa "
        "— mtu anayeelewa nchi yake, si tu kuiita jina.",
    ),
    (
        "Unatarajia kuwafikia vijana wangapi?",
        "Mwaka wa kwanza (realistic): takriban 5,000–10,000 vijana — kupitia tovuti, mitandao ya "
        "kijamii, na ushirikiano na shule/vyuo 5–10.\n\n"
        "Mwaka wa pili (ukiwa na ushirikiano): 50,000+ — ikiwa tutaunganishwa na programu za TEHAMA "
        "za kitaifa (kama Girls in ICT), vilabu vya coding, na mtaala wa Historia/Uraia.\n\n"
        "Tunapima mafanikio kwa: idadi ya watumiaji, maswali yaliyoulizwa, na muda wanaotumia "
        "kujifunza — si tu “views” za mitandao.\n\n"
        "Lengo la muda mrefu: kufanya MuunganoAi iwe msaidizi wa kawaida wa kujifunza kuhusu "
        "Muungano wa Tanzania — kama Wikipedia ya Muungano, lakini ya mazungumzo na ya vyanzo rasmi.",
    ),
]


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x2E)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def add_form_answers(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("MAJIBU YA FOMU — NAKILI MOJA KWA MOJA")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1A, 0x5C, 0x2E)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(
        "Girls in ICT AI Hackathon 2026  |  MuunganoAi  |  "
        "Chagua kila jibu chini ya swali na nakili kwenye fomu"
    )
    nr.italic = True
    nr.font.size = Pt(9)

    doc.add_paragraph()

    for idx, (question, answer) in enumerate(FORM_ANSWERS, start=1):
        q = doc.add_paragraph()
        qr = q.add_run(f"{idx}. {question}")
        qr.bold = True
        qr.font.size = Pt(12)
        qr.font.color.rgb = RGBColor(0x1A, 0x5C, 0x2E)
        q.paragraph_format.space_before = Pt(10)
        q.paragraph_format.space_after = Pt(4)

        for block in answer.split("\n\n"):
            if block.strip():
                add_para(doc, block.strip())
        doc.add_paragraph()


def add_proposal_content(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PROPOSAL YA MRADI")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1A, 0x5C, 0x2E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("MuunganoAi")
    r2.bold = True
    r2.font.size = Pt(16)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = sub2.add_run("Msaidizi wa AI wa Elimu ya Muungano wa Tanganyika na Zanzibar")
    r3.italic = True
    r3.font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = meta.add_run(
        "Mada: Akili Unde kwa Maendeleo — Elimu ya Uraia kwa Kizazi cha Kidijitali\n"
        "Toleo: 1.0  |  Julai 2026  |  Hali: ~90% imekamilika"
    )
    m.font.size = Pt(10)

    doc.add_paragraph()

    add_heading(doc, "1. Muhtasari wa Mkurugenzi", 1)
    add_para(
        doc,
        "MuunganoAi ni mfumo wa akili bandia unaosaidia Watanzania — hasa vijana — kuelewa "
        "Muungano wa Tanganyika na Zanzibar kwa njia rahisi, ya kuaminika, na inayopatikana. "
        "Badala ya kusoma mamia ya kurasa za Katiba na Hati za Muungano, mtumiaji anauliza swali "
        "na kupata jibu linalotokana na nyaraka rasmi tu — si kutoka kwa mawazo ya AI peke yake.",
    )
    add_para(
        doc,
        "Mfumo unatumia teknolojia ya RAG (Retrieval-Augmented Generation): kwanza unatafuta taarifa "
        "sahihi kutoka maktaba ya nyaraka zilizothibitishwa, kisha AI inaandika jibu kulingana na "
        "vyanzo hivyo. Mtumiaji anaweza kuwasiliana kwa maandishi au sauti, kwa Kiswahili au "
        "Kiingereza, na kwa mtindo unaomfaa (vijana, mwanafunzi, rasmi).",
    )
    add_para(doc, "Lengo kuu: kuongeza uelewa wa Muungano, si kuonyesha teknolojia kwa ajili ya teknolojia.", True)

    add_heading(doc, "2. Mandhari na Tatizo", 1)
    add_para(doc, "Tanzania ina historia na muundo wa kipekee — Muungano wa nchi mbili chini ya bendera moja. Maswali muhimu bado yana changamoto kwa umma:")
    add_bullets(
        doc,
        [
            "Muungano ulianzishwa lini na kwa namna gani?",
            "Kwa nini kuna serikali mbili?",
            '"Mambo ya Muungano" (Union Matters) ni nini hasa?',
            "Tofauti kati ya mamlaka ya Bara na Zanzibar ni ipi?",
        ],
    )
    add_table(
        doc,
        ["Tatizo", "Athari"],
        [
            ["Nyaraka rasmi ni ndefu na ngumu", "Vijana na raia wengi hawasomi Katiba kwa undani"],
            ["Taarifa mtandaoni ni mchanganyiko", "Ukweli, maoni, na uvumi vimechanganyika"],
            ["Hakuna zana ya elimu inayopatikana", "Elimu ya uraia inabaki darasani tu"],
            ["Lugha na mitindo tofauti", "Mfumo mmoja haufikii watumiaji wote"],
        ],
    )

    add_heading(doc, "3. Suluhisho: MuunganoAi", 1)
    add_para(
        doc,
        'MuunganoAi ni msaidizi wa mazungumzo maalum wa elimu ya Muungano. Si chatbot ya jumla — '
        'ni zana iliyojengwa kwa ajili ya swali moja kuu: "Nifundishe Muungano wangu."',
    )
    add_para(doc, "Mtiririko wa kazi:", True)
    add_bullets(
        doc,
        [
            "Mtumiaji anauliza (maandishi au sauti)",
            "Mfumo unatafuta vipande vinavyohusiana (ChromaDB + embeddings za lugha mbili)",
            "AI inaandika jibu kutoka muktadha uliopatikana (Gemini 2.5 Flash)",
            "Mtumiaji anapokea jibu la mazungumzo + marejeo ya vyanzo (mtindo rasmi)",
        ],
    )
    add_para(doc, "Sifa kuu:", True)
    add_bullets(
        doc,
        [
            "Vyanzo vilivyothibitishwa — allowlist ya tovuti na nyaraka rasmi",
            "Kipaumbele cha nyaraka — Katiba na Articles of Union kwanza",
            "Lugha mbili — Kiswahili na Kiingereza",
            "Mitindo 5 ya majibu — youth, simple, student, official, default",
            "Mazungumzo ya kawaida — andika au zungumza; soma au sikiliza jibu",
            "Streaming — majibu yanakuja moja kwa moja",
        ],
    )

    add_heading(doc, "4. Teknolojia", 1)
    add_table(
        doc,
        ["Tabaka", "Teknolojia", "Kazi"],
        [
            ["Backend", "Python 3.10+, FastAPI", "API, chat, ingest"],
            ["Hifadhi ya maarifa", "ChromaDB", "Vector store ya vipande vya nyaraka"],
            ["Embeddings", "paraphrase-multilingual-MiniLM-L12-v2", "Utafutaji wa semantic (SW/EN)"],
            ["LLM", "Gemini 2.5 Flash (OpenRouter)", "Majibu ya mazungumzo"],
            ["Crawler", "Trafilatura + allowlist YAML", "Kupakia vyanzo rasmi"],
            ["Frontend", "HTML, CSS, JavaScript", "UI ya mazungumzo"],
            ["Sauti", "Web Speech API", "STT/TTS — sehemu ya kawaida ya UX"],
        ],
    )
    add_para(
        doc,
        "Kwa nini RAG? Mfumo wa kawaida wa AI unaweza kubuni tarehe na ukweli. RAG inahakikisha "
        "kila jibu lina msingi wa nyaraka halisi — kanuni muhimu kwa elimu ya uraia.",
    )

    add_heading(doc, "5. Lengo na Watumiaji", 1)
    add_para(doc, "Watumiaji wa msingi:", True)
    add_bullets(
        doc,
        [
            "Vijana (shule ya sekondari na vyuo)",
            "Wanafunzi wa Historia, Siasa, na Sheria",
            "Walimu wanaohitaji msaidizi wa kufafanua mambo magumu",
            "Raia wanaotaka kuelewa nchi yao bila kusoma PDF ndefu",
        ],
    )

    add_heading(doc, "6. Jinsi Tutakavyowafikia Vijana", 1)
    add_bullets(
        doc,
        [
            "Tovuti inayopatikana kwenye simu — hakuna usakinishaji",
            "Mitandao ya kijamii — maswali mafupi + majibu (Instagram, TikTok, WhatsApp)",
            "Shule na vyuo — msaidizi wa walimu wa Historia na Uraia",
            "Vilabu vya TEHAMA — kuonyesha AI kwa elimu ya uraia",
        ],
    )

    add_heading(doc, "7. Athari Inayotarajiwa", 1)
    add_para(doc, "Matokeo yanayotarajiwa:", True)
    add_bullets(
        doc,
        [
            "Vijana wanaelewa historia na muundo wa nchi yao",
            "Kupungua kwa uvumi na taarifa potofu mtandaoni",
            "Kuongezeka kwa mazungumzo ya kiufundi kuhusu Muungano",
            "Msingi bora wa uraia wenye maarifa",
        ],
    )
    add_table(
        doc,
        ["SDG", "Uchanganuzi"],
        [
            ["4 — Elimu bora", "Elimu ya uraia inayopatikana, lugha mbili"],
            ["16 — Amani na haki", "Uelewa wa taasisi za kitaifa"],
            ["5 — Usawa wa kijinsia", "Wasichana katika TEHAMA wanajenga AI ya kuaminika"],
        ],
    )
    add_table(
        doc,
        ["Kipindi", "Lengo la Ufikiaji"],
        [
            ["Mwaka 1", "5,000 – 10,000 vijana"],
            ["Mwaka 2", "50,000+ (ushirikiano na programu za TEHAMA)"],
            ["Muda mrefu", "Msaidizi wa kawaida wa kujifunza kuhusu Muungano"],
        ],
    )

    add_heading(doc, "8. Hali ya Sasa na Mpango", 1)
    add_para(doc, "Imekamilika (~90%):", True)
    add_bullets(
        doc,
        [
            "Mfumo wa RAG kamili",
            "Maktaba ya maarifa: 600+ vipande vya nyaraka rasmi",
            "Chat, streaming, lugha mbili, mitindo 5",
            "UI ya mazungumzo (simu + desktop)",
            "Uwezo wa sauti (ingizo na majibu)",
        ],
    )
    add_para(doc, "Inayobaki (~10%): uboreshaji wa UI, upanuzi wa vyanzo, pilot shuleni, hosting ya umma.", True)

    add_heading(doc, "9. Uaminifu na Kanuni", 1)
    add_bullets(
        doc,
        [
            "Majibu yanategemea vyanzo vilivyoidhinishwa",
            "Tarehe na ukweli maalum — kamwe si makisio ya modeli",
            "Vyanzo visivyothibitishwa vimezuiwa",
            "Zana ya elimu — si ushauri wa kisheria, si mwakilishi wa serikali",
        ],
    )

    add_heading(doc, "10. Timu", 1)
    add_table(
        doc,
        ["Jina", "Wajibu", "Taasisi"],
        [
            ["[Jaza jina]", "[Lead / Developer]", "[Chuo]"],
            ["[Jaza jina]", "[Mhusika]", "[Chuo]"],
        ],
    )

    add_heading(doc, "11. Mahitaji ya Msaada", 1)
    add_table(
        doc,
        ["Kipengele", "Maelezo"],
        [
            ["Mentorship", "AI ethics, RAG best practices"],
            ["Ushirikiano", "Wizara ya Elimu, TCRA, vilabu vya TEHAMA"],
            ["Hosting", "Server ya umma"],
        ],
    )

    add_heading(doc, "12. Hitimisho", 1)
    add_para(
        doc,
        "MuunganoAi inaonyesha jinsi akili bandia inavyoweza kutumika kwa maendeleo halisi — "
        "elimu ya uraia inayopatikana kwa kizazi kinachokua na kidijitali.",
    )
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = closing.add_run(
        "MuunganoAi — elewa Muungano wako, kwa lugha yako, kutoka vyanzo unavyoweza kuamini."
    )
    cr.bold = True
    cr.italic = True
    cr.font.size = Pt(12)

    doc.add_paragraph()
    add_para(doc, "Mawasiliano: [Jaza: barua pepe, simu, GitHub, demo URL]", True)
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(
        'Hati hii ni sehemu ya ombi la Girls in ICT AI Hackathon 2026 — '
        '"Akili Unde kwa Maendeleo: Wasichana Wakibadilisha Mustakabali wa Kidijitali."'
    )
    fr.italic = True
    fr.font.size = Pt(9)


def new_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    return doc


def build_proposal() -> None:
    doc = new_document()
    add_proposal_content(doc)
    doc.save(OUT_PROPOSAL)
    print(f"Created: {OUT_PROPOSAL}")


def build_submission() -> None:
    doc = new_document()
    add_form_answers(doc)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_proposal_content(doc)
    doc.save(OUT_SUBMISSION)
    print(f"Created: {OUT_SUBMISSION}")


def build() -> None:
    build_proposal()
    build_submission()


if __name__ == "__main__":
    build()
